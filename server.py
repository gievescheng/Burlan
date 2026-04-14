from __future__ import annotations

import json
import os
import re
import secrets
import tempfile
import time
import traceback
import zipfile
import io
from pathlib import Path
from urllib.parse import urlencode
from datetime import datetime

import requests
from docx import Document
from openpyxl import load_workbook
from flask import (
    Flask,
    Response,
    abort,
    jsonify,
    redirect,
    request,
    send_file,
    send_from_directory,
    session,
)
import ops_data
import burlan_document_sources
import server_handlers_documents
import server_handlers_integrations
import server_handlers_operations
from burlan_paths import (
    BURLAN_AUDIT_DIR,
    BURLAN_CALIBRATION_DIR,
    BURLAN_EQUIPMENT_DIR,
    BURLAN_MASTER_LIST_PATH,
    BURLAN_OBJECTIVE_DIR,
    BURLAN_SUPPLIER_DIR,
    BURLAN_THIRD_LEVEL_DIR,
)
from runtime_paths import (
    GOOGLE_CONFIG_PATH,
    GOOGLE_TOKEN_PATH,
    get_or_create_flask_secret,
    migrate_legacy_private_files,
    public_root_contains_private_files,
)

BASE_DIR = Path(__file__).parent.resolve()
GOOGLE_AUTH_URL = 'https://accounts.google.com/o/oauth2/v2/auth'
GOOGLE_TOKEN_URL = 'https://oauth2.googleapis.com/token'
GOOGLE_USERINFO_URL = 'https://www.googleapis.com/oauth2/v2/userinfo'
GOOGLE_CALENDAR_EVENTS_URL = 'https://www.googleapis.com/calendar/v3/calendars/primary/events'
GOOGLE_SCOPES = [
    'https://www.googleapis.com/auth/calendar.events',
    'https://www.googleapis.com/auth/userinfo.email',
    'https://www.googleapis.com/auth/userinfo.profile',
]
PUBLIC_STATIC_DIRS = {
    'vendor': BASE_DIR / 'vendor',
    'public': BASE_DIR / 'public',
}
PUBLIC_STATIC_FILES = {
    'index.html': BASE_DIR / 'index.html',
}
LEGACY_EQUIPMENT_IDS = {'JE-001', 'JE-002', 'JE-003', 'JE-004', 'JE-005', 'JE-006'}
LEGACY_SUPPLIER_IDS = {'SUP-001', 'SUP-002', 'SUP-003', 'SUP-004', 'SUP-005', 'SUP-006'}
LEGACY_SUPPLIER_NAMES = {
    '楊特企業有限公司',
    '金華瑋科技有限公司',
    '柏連企業股份有限公司',
    '鏵友益科技股份有限公司',
    '奈米趨勢科技有限公司',
    '拓生科技有限公司',
}

app = Flask(__name__, static_folder=None)
app.config['MAX_CONTENT_LENGTH'] = 64 * 1024 * 1024
migrate_legacy_private_files()
app.secret_key = os.environ.get('FLASK_SECRET_KEY') or get_or_create_flask_secret()


def read_json(path: Path, fallback):
    if not path.exists():
        return fallback
    try:
        return json.loads(path.read_text(encoding='utf-8'))
    except Exception:
        return fallback


def write_json(path: Path, data) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')


def delete_file(path: Path) -> None:
    if path.exists():
        path.unlink()


def load_google_config() -> dict:
    return read_json(GOOGLE_CONFIG_PATH, {})


def save_google_config(client_id: str, client_secret: str) -> dict:
    data = {
        'client_id': client_id,
        'client_secret': client_secret,
        'updated_at': int(time.time()),
    }
    write_json(GOOGLE_CONFIG_PATH, data)
    return data


def clear_google_config() -> None:
    delete_file(GOOGLE_CONFIG_PATH)


def load_google_tokens() -> dict:
    return read_json(GOOGLE_TOKEN_PATH, {})


def save_google_tokens(tokens: dict) -> dict:
    tokens = dict(tokens)
    tokens['updated_at'] = int(time.time())
    write_json(GOOGLE_TOKEN_PATH, tokens)
    return tokens


def clear_google_tokens() -> None:
    delete_file(GOOGLE_TOKEN_PATH)


def google_redirect_uri() -> str:
    return request.host_url.rstrip('/') + '/api/google-calendar/oauth/callback'


def google_configured(config: dict | None = None) -> bool:
    config = config or load_google_config()
    return bool(config.get('client_id') and config.get('client_secret'))


def enrich_google_identity(tokens: dict) -> dict:
    access_token = tokens.get('access_token')
    if not access_token:
        return tokens
    try:
        response = requests.get(
            GOOGLE_USERINFO_URL,
            headers={'Authorization': f'Bearer {access_token}'},
            timeout=20,
        )
        if response.ok:
            userinfo = response.json()
            tokens['email'] = userinfo.get('email', tokens.get('email', ''))
            tokens['name'] = userinfo.get('name', tokens.get('name', ''))
    except Exception:
        pass
    return tokens


def exchange_google_code(code: str, config: dict) -> dict:
    response = requests.post(
        GOOGLE_TOKEN_URL,
        data={
            'code': code,
            'client_id': config['client_id'],
            'client_secret': config['client_secret'],
            'redirect_uri': google_redirect_uri(),
            'grant_type': 'authorization_code',
        },
        timeout=20,
    )
    if not response.ok:
        raise RuntimeError(google_error_text(response))
    tokens = response.json()
    tokens['expires_at'] = int(time.time()) + int(tokens.get('expires_in', 3600)) - 60
    return enrich_google_identity(tokens)


def refresh_google_tokens(tokens: dict, config: dict) -> dict:
    expires_at = int(tokens.get('expires_at', 0) or 0)
    if expires_at > int(time.time()) + 30:
        return tokens

    refresh_token = tokens.get('refresh_token')
    if not refresh_token:
        raise RuntimeError('Google Calendar authorization expired. Please reconnect Google Calendar.')

    response = requests.post(
        GOOGLE_TOKEN_URL,
        data={
            'client_id': config['client_id'],
            'client_secret': config['client_secret'],
            'refresh_token': refresh_token,
            'grant_type': 'refresh_token',
        },
        timeout=20,
    )
    if not response.ok:
        raise RuntimeError(google_error_text(response))

    refreshed = response.json()
    merged = dict(tokens)
    merged.update(refreshed)
    merged['refresh_token'] = refresh_token
    merged['expires_at'] = int(time.time()) + int(refreshed.get('expires_in', 3600)) - 60
    return save_google_tokens(enrich_google_identity(merged))


def require_google_access_token() -> tuple[str, dict]:
    config = load_google_config()
    if not google_configured(config):
        raise RuntimeError('Google Calendar is not configured. Fill in Client ID and Client Secret first.')

    tokens = load_google_tokens()
    if not tokens.get('access_token'):
        raise RuntimeError('Google Calendar is not connected. Complete Google authorization first.')

    tokens = refresh_google_tokens(tokens, config)
    return tokens['access_token'], tokens


def google_error_text(response: requests.Response) -> str:
    try:
        payload = response.json()
        if isinstance(payload.get('error'), dict):
            return payload['error'].get('message', json.dumps(payload, ensure_ascii=False))
        return payload.get('error_description') or payload.get('error') or json.dumps(payload, ensure_ascii=False)
    except Exception:
        return response.text.strip() or f'HTTP {response.status_code}'


def json_error(message: str, status_code: int = 400):
    return jsonify({'error': message}), status_code


def json_body() -> dict:
    return request.get_json(force=True) or {}


def load_burlan_master_documents() -> dict:
    original_path = burlan_document_sources.BURLAN_MASTER_LIST_PATH
    burlan_document_sources.BURLAN_MASTER_LIST_PATH = BURLAN_MASTER_LIST_PATH
    try:
        return burlan_document_sources.load_burlan_master_documents()
    finally:
        burlan_document_sources.BURLAN_MASTER_LIST_PATH = original_path


def normalize_document_title(value: str) -> str:
    return burlan_document_sources.normalize_document_title(value)


def build_burlan_document_name_lookup(documents: list[dict]) -> dict[str, str]:
    return burlan_document_sources.build_burlan_document_name_lookup(documents)


def parse_burlan_schedule_date(text: str) -> str:
    source = str(text or '').strip()
    match = re.search(r'自\s*(\d{2,4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', source)
    if not match:
        return ''
    year = int(match.group(1))
    if year < 1911:
        year += 1911
    return ops_data._parse_date(f'{year}/{match.group(2)}/{match.group(3)}')


def preserve_unique(values: list[str]) -> list[str]:
    seen = set()
    result = []
    for value in values:
        text = str(value or '').strip()
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)
    return result


def collect_burlan_audit_attachments(year_dir: Path, plan_file: Path) -> list[str]:
    attachments: list[str] = []

    def append_if_exists(path: Path | None):
        if path is None or not path.exists():
            return
        resolved = str(path.resolve())
        if resolved not in attachments:
            attachments.append(resolved)

    append_if_exists(plan_file)
    append_if_exists(next((path for path in year_dir.iterdir() if path.is_file() and path.name.lower().startswith('17.1') and path.suffix.lower() == '.pdf'), None))

    for prefix in ('17.4', '17.5'):
        for suffix in ('.docx', '.doc', '.pdf'):
            candidate = next((path for path in year_dir.iterdir() if path.is_file() and path.name.lower().startswith(prefix.lower()) and path.suffix.lower() == suffix), None)
            if candidate:
                append_if_exists(candidate)
                break

    checklist_dir = next((path for path in year_dir.iterdir() if path.is_dir() and path.name.startswith('17.3')), None)
    if checklist_dir:
        sample_checklists = sorted([path for path in checklist_dir.iterdir() if path.is_file() and path.suffix.lower() in {'.docx', '.doc', '.pdf'}])[:3]
        for path in sample_checklists:
            append_if_exists(path)

    return attachments


def build_burlan_annual_audit_plan(plan_file: Path, document_lookup: dict[str, str]) -> dict | None:
    rows = ops_data._docx_row_values(plan_file)
    if len(rows) < 4:
        return None

    scheduled_date = parse_burlan_schedule_date(' '.join(rows[0]))
    if not scheduled_date:
        return None

    detail_rows = []
    for row in rows[3:]:
        if len(row) < 6:
            continue
        no_text = str(row[0] or '').strip()
        if not no_text.isdigit():
            continue
        dept = str(row[1] or '').strip()
        auditee = str(row[2] or '').strip()
        auditor = str(row[3] or '').strip()
        scope_name = str(row[5] or '').strip()
        if not any((dept, auditee, auditor, scope_name)):
            continue
        detail_rows.append({
            'dept': dept,
            'auditee': auditee,
            'auditor': auditor,
            'scope_name': scope_name,
        })

    if not detail_rows:
        return None

    schedule_year = int(scheduled_date[:4])
    scope_codes = preserve_unique(
        document_lookup.get(normalize_document_title(item['scope_name']), '')
        for item in detail_rows
    )
    scope_names = preserve_unique(item['scope_name'] for item in detail_rows)
    attachments = collect_burlan_audit_attachments(plan_file.parent, plan_file)
    status = '已完成' if scheduled_date <= datetime.now().date().isoformat() else '計畫中'

    return {
        'id': f'IA-{schedule_year}-01',
        'year': schedule_year,
        'period': '年度',
        'scheduledDate': scheduled_date,
        'dept': '、'.join(preserve_unique(item['dept'] for item in detail_rows)),
        'scope': ','.join(scope_codes) if scope_codes else '、'.join(scope_names),
        'auditor': '、'.join(preserve_unique(item['auditor'] for item in detail_rows)),
        'auditee': '、'.join(preserve_unique(item['auditee'] for item in detail_rows)),
        'status': status,
        'actualDate': scheduled_date if status == '已完成' else '',
        'findings': 0,
        'ncCount': 0,
        'source_file': str(plan_file.resolve()),
        'attachment_paths': attachments,
        'source_system': 'burlan_audit_plan',
        'read_only': True,
    }


def merge_burlan_audit_plan_with_system(source_plan: dict, system_plan: dict) -> dict:
    merged = dict(source_plan)
    for field in ('actualDate', 'status', 'findings', 'ncCount', 'created_at', 'updated_at'):
        value = system_plan.get(field)
        if value not in (None, ''):
            merged[field] = value
    for field in ('dept', 'scope', 'auditor', 'auditee'):
        if not merged.get(field) and system_plan.get(field):
            merged[field] = system_plan[field]
    extra_attachments = [str(path).strip() for path in (system_plan.get('attachment_paths') or []) if str(path).strip()]
    merged['attachment_paths'] = preserve_unique(list(merged.get('attachment_paths') or []) + extra_attachments)
    return merged


def load_burlan_audit_plans() -> dict:
    if not BURLAN_AUDIT_DIR.exists():
        items = ops_data.load_records('auditplan')
        return {
            'items': items,
            'count': len(items),
            'source_path': str(BURLAN_AUDIT_DIR),
            'message': '找不到 17內部稽核管理程序 記錄資料夾，已改用系統資料。',
            'source_system': 'system_auditplan',
            'mode': 'fallback',
        }

    master_documents = load_burlan_master_documents().get('items', [])
    document_lookup = build_burlan_document_name_lookup(master_documents)
    source_items = []
    for year_dir in sorted([path for path in BURLAN_AUDIT_DIR.iterdir() if path.is_dir()], key=lambda item: item.name):
        plan_file = next((path for path in year_dir.iterdir() if path.is_file() and path.name.lower().startswith('17.1') and path.suffix.lower() == '.docx'), None)
        if not plan_file:
            continue
        parsed = build_burlan_annual_audit_plan(plan_file, document_lookup)
        if parsed:
            source_items.append(parsed)

    if not source_items:
        items = ops_data.load_records('auditplan')
        return {
            'items': items,
            'count': len(items),
            'source_path': str(BURLAN_AUDIT_DIR),
            'message': '尚未讀到柏連年度內部稽核計畫表，已改用系統資料。',
            'source_system': 'system_auditplan',
            'mode': 'fallback',
        }

    stored_items = ops_data.load_records('auditplan')
    source_by_id = {item['id']: item for item in source_items}
    source_years = {int(item.get('year') or 0) for item in source_items}

    for item in stored_items:
        item_id = str(item.get('id') or '').strip()
        if item_id in source_by_id:
            source_by_id[item_id] = merge_burlan_audit_plan_with_system(source_by_id[item_id], item)

    merged_items = list(source_by_id.values())
    for item in stored_items:
        item_id = str(item.get('id') or '').strip()
        if item_id in source_by_id:
            continue
        stored_year = int(item.get('year') or 0)
        has_source_file = bool(str(item.get('source_file') or '').strip())
        if stored_year in source_years and not has_source_file:
            continue
        merged_items.append(dict(item, source_system='system_auditplan', read_only=False))

    merged_items.sort(key=lambda item: (str(item.get('scheduledDate') or ''), str(item.get('id') or '')), reverse=True)
    return {
        'items': merged_items,
        'count': len(merged_items),
        'source_path': str(BURLAN_AUDIT_DIR),
        'message': f'已載入柏連年度內部稽核計畫 {len(source_items)} 筆，並合併系統自建資料。',
        'source_system': 'burlan_audit_plan',
        'mode': 'source',
    }


def get_audit_plan_source_year(plan: dict) -> str:
    source_file = str(plan.get('source_file') or '').strip()
    if not source_file:
        return ''
    for part in Path(source_file).parts:
        if str(part).endswith('年度'):
            return str(part)
    return ''


def normalize_audit_year_label(value: str) -> str:
    text = str(value or '').strip()
    match = re.search(r'(\d{2,4})', text)
    if not match:
        return text
    return f"{match.group(1)}年度"


def build_audit_plan_year_bundle(year_filter: str) -> tuple[io.BytesIO, str]:
    payload = load_burlan_audit_plans()
    items = payload.get('items', [])
    normalized_year = normalize_audit_year_label(year_filter)
    filtered_items = [
        item for item in items
        if normalize_audit_year_label(get_audit_plan_source_year(item)) == normalized_year
    ]
    if not filtered_items:
        raise ValueError(f'找不到 {normalized_year or year_filter} 的稽核計畫資料。')

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', compression=zipfile.ZIP_DEFLATED) as archive:
        manifest_lines = [f'柏連內部稽核附件包：{normalized_year or year_filter}', '']
        for item in filtered_items:
            manifest_lines.append(f"{item.get('id', '')}｜{item.get('dept', '')}｜{item.get('scheduledDate', '')}")
            attachments = ops_data.list_auditplan_attachments(str(item.get('id') or ''), items)
            if not attachments:
                manifest_lines.append('  - 無可下載附件')
                continue
            for attachment in attachments:
                if not attachment.get('exists'):
                    manifest_lines.append(f"  - 缺檔：{attachment.get('name') or attachment.get('path')}")
                    continue
                stored_path = str(attachment.get('path') or '')
                file_path = ops_data.get_serving_path(stored_path)
                if file_path is None or not file_path.exists():
                    manifest_lines.append(f"  - 缺檔：{attachment.get('name') or stored_path}")
                    continue
                arcname = f"{normalized_year or year_filter}/{item.get('id', '未命名計畫')}/{file_path.name}"
                archive.write(file_path, arcname)
                manifest_lines.append(f"  - {file_path.name}")
        archive.writestr(f'{normalized_year or year_filter}/README.txt', '\n'.join(manifest_lines))

    buffer.seek(0)
    return buffer, f'柏連內部稽核附件包_{normalized_year or year_filter}.zip'


def normalize_objective_month_value(value: str) -> str:
    return re.sub(r'\s+', ' ', str(value or '').strip())


def month_value_status(value: str) -> str:
    text = normalize_objective_month_value(value)
    if not text:
        return 'pending'
    lowered = text.lower()
    if any(keyword in text for keyword in ('未達', '不符合', '異常', '不足', '逾期', '不合格')):
        return 'missed'
    if any(keyword in text for keyword in ('符合', '達', '合格', '完成', '正常', '穩定')):
        return 'achieved'
    if any(keyword in lowered for keyword in ('pass', 'ok')):
        return 'achieved'
    if any(keyword in lowered for keyword in ('fail', 'ng')):
        return 'missed'
    return 'manual_review'


def pick_latest_objective_file() -> Path | None:
    if not BURLAN_OBJECTIVE_DIR.exists():
        return None
    files = sorted(BURLAN_OBJECTIVE_DIR.glob('3.1品質目標管制表(*年度).docx'))
    if not files:
        return None

    def file_year(path: Path) -> int:
        match = re.search(r'\((\d+)年度\)', path.stem)
        return int(match.group(1)) if match else 0

    return sorted(files, key=file_year)[-1]


def load_burlan_quality_objectives() -> dict:
    source_file = pick_latest_objective_file()
    if source_file is None:
        return {
            'items': [],
            'count': 0,
            'achieved_count': 0,
            'pending_count': 0,
            'manual_review_count': 0,
            'achievement_rate': None,
            'source_system': 'burlan_quality_objectives',
            'source_path': str(BURLAN_OBJECTIVE_DIR),
            'message': '3 目標管理程序記錄中尚未找到可讀取的品質目標管制表。',
        }

    document = Document(source_file)
    if not document.tables:
        return {
            'items': [],
            'count': 0,
            'achieved_count': 0,
            'pending_count': 0,
            'manual_review_count': 0,
            'achievement_rate': None,
            'source_system': 'burlan_quality_objectives',
            'source_path': str(source_file),
            'message': '品質目標管制表沒有讀到可用表格。',
        }

    table = document.tables[0]
    month_labels = [str(cell.text or '').strip() for cell in table.rows[1].cells[7:19]]
    objective_map = {}

    for row in table.rows[2:]:
        cells = [str(cell.text or '').strip() for cell in row.cells]
        if len(cells) < 19:
            continue
        item_no = cells[0].strip()
        objective_name = cells[2].strip()
        row_type = re.sub(r'\s+', '', cells[6])
        if not item_no or not objective_name or row_type not in {'目標', '實績', '判定'}:
            continue

        entry = objective_map.setdefault(item_no, {
            'id': f'QO-{item_no}',
            'item_no': item_no,
            'quality_policy': cells[1].strip(),
            'objective': objective_name,
            'department': cells[3].strip(),
            'resource': cells[4].strip(),
            'measurement': cells[5].strip(),
            'target_by_month': {},
            'actual_by_month': {},
            'judgement_by_month': {},
            'remark': cells[19].strip() if len(cells) > 19 else '',
        })
        month_values = cells[7:19]
        target_key = {
            '目標': 'target_by_month',
            '實績': 'actual_by_month',
            '判定': 'judgement_by_month',
        }[row_type]
        month_bucket = entry[target_key]
        for month_label, raw_value in zip(month_labels, month_values):
            month = str(month_label or '').strip()
            if not month:
                continue
            value = normalize_objective_month_value(raw_value)
            if value:
                month_bucket[month] = value

    items = []
    achieved_count = 0
    pending_count = 0
    manual_review_count = 0

    for item_no in sorted(objective_map, key=lambda value: int(re.sub(r'\D', '', value) or '0')):
        entry = objective_map[item_no]
        latest_month = ''
        latest_target = ''
        latest_actual = ''
        latest_judgement = ''
        for month_label in reversed(month_labels):
            month = str(month_label or '').strip()
            target_value = entry['target_by_month'].get(month, '')
            actual_value = entry['actual_by_month'].get(month, '')
            judgement_value = entry['judgement_by_month'].get(month, '')
            if target_value or actual_value or judgement_value:
                latest_month = month
                latest_target = target_value
                latest_actual = actual_value
                latest_judgement = judgement_value
                break

        status = month_value_status(latest_judgement or latest_actual)
        if status == 'achieved':
            achieved_count += 1
        elif status == 'pending':
            pending_count += 1
        else:
            manual_review_count += 1

        items.append({
            **entry,
            'latest_month': latest_month,
            'latest_target': latest_target,
            'latest_actual': latest_actual,
            'latest_judgement': latest_judgement,
            'status': status,
        })

    total = len(items)
    achievement_rate = round(achieved_count / total * 100) if total > 0 else None
    file_year_match = re.search(r'\((\d+)年度\)', source_file.stem)
    file_year = file_year_match.group(1) if file_year_match else ''
    return {
        'items': items,
        'count': total,
        'achieved_count': achieved_count,
        'pending_count': pending_count,
        'manual_review_count': manual_review_count,
        'achievement_rate': achievement_rate,
        'source_system': 'burlan_quality_objectives',
        'source_path': str(source_file),
        'year': file_year,
        'message': f'已讀取品質目標管制表（{file_year}年度）共 {total} 項目標。',
    }


def _clean_docx_text(value: str) -> str:
    return str(value or '').replace('\xa0', ' ').strip()


def _canonical_label(value: str) -> str:
    return re.sub(r'[\s:：]+', '', _clean_docx_text(value))


def _normalize_iso_date(value: str) -> str:
    text = _clean_docx_text(value).replace('/', '.').replace('-', '.')
    match = re.search(r'(\d{3,4})\.(\d{1,2})\.(\d{1,2})', text)
    if not match:
        return ''
    year = int(match.group(1))
    month = int(match.group(2))
    day = int(match.group(3))
    if year < 1911:
        year += 1911
    try:
        return datetime(year, month, day).date().isoformat()
    except ValueError:
        return ''


def _frequency_to_days(value: str) -> int | None:
    text = _clean_docx_text(value)
    if not text or text.upper() == 'NA' or '免校' in text:
        return None
    if text in {'每年', '年度', '一年', '每1年'} or ('每' in text and '年' in text):
        return 365
    if '半年' in text:
        return 183
    year_match = re.search(r'(\d+)\s*年', text)
    if year_match:
        return int(year_match.group(1)) * 365
    month_match = re.search(r'(\d+)\s*個?月', text)
    if month_match:
        return int(month_match.group(1)) * 30
    if '月' in text:
        return 30
    if '季' in text:
        return 90
    return None


def _guess_instrument_type(name: str) -> str:
    text = _clean_docx_text(name).upper()
    if 'FTIR' in text:
        return 'FTIR'
    if 'PH' in text:
        return 'PH'
    if 'COD' in text:
        return 'COD'
    if '黏度' in text:
        return '黏度'
    if '電導' in text:
        return '電導'
    if '折射' in text:
        return '折射'
    if '水浴' in text:
        return '水浴鍋'
    if '測試筆' in text:
        return '測試筆'
    if '電子台秤' in text or '桶秤' in text or '充填秤' in text:
        return '秤重設備'
    return _clean_docx_text(name)[:12]


def _extract_instrument_alias(path: Path) -> str:
    stem = path.stem
    stem = re.sub(r'^9\.2\.\d+量規儀器履歷表', '', stem)
    return _clean_docx_text(stem)


def _extract_docx_pairs(path: Path) -> dict[str, str]:
    doc = Document(path)
    pairs: dict[str, str] = {}
    if not doc.tables:
        return pairs
    for row in doc.tables[0].rows:
        cells = [_clean_docx_text(cell.text) for cell in row.cells]
        for idx in range(0, len(cells) - 1, 2):
            label = _canonical_label(cells[idx])
            value = cells[idx + 1]
            if label and value and label not in pairs:
                pairs[label] = value
    return pairs


def _cell_checked(value: str) -> bool:
    text = _clean_docx_text(value).replace(' ', '')
    return text not in {'', '-', '--', 'NA', 'N/A'}


def _split_keeper_text(value: str) -> tuple[str, str]:
    text = _clean_docx_text(value)
    if '/' in text:
        left, right = [part.strip() for part in text.split('/', 1)]
        return left, right
    return '', text


def _load_calibration_inventory(records_dir: Path) -> tuple[dict[str, dict[str, str]], Path | None]:
    inventory_path = next(records_dir.glob('9.1*.docx'), None)
    if inventory_path is None:
        return {}, None
    doc = Document(inventory_path)
    if not doc.tables:
        return {}, inventory_path

    inventory: dict[str, dict[str, str]] = {}
    table = doc.tables[0]
    for row in table.rows[2:]:
        cells = [_clean_docx_text(cell.text) for cell in row.cells]
        if len(cells) < 13:
            continue
        instrument_code = cells[0]
        instrument_name = cells[1]
        if not instrument_code or not instrument_name:
            continue
        location, keeper = _split_keeper_text(cells[11])
        calib_method = ''
        if _cell_checked(cells[10]):
            calib_method = '免校'
        elif _cell_checked(cells[8]):
            calib_method = '內校'
        elif _cell_checked(cells[9]):
            calib_method = '外校'
        inventory[instrument_code] = {
            'id': instrument_code,
            'name': instrument_name,
            'spec': cells[2],
            'model': cells[3],
            'serialNo': cells[4],
            'brand': cells[5],
            'tolerance': cells[6],
            'frequencyLabel': cells[7],
            'calibMethod': calib_method,
            'location': location,
            'keeper': keeper,
            'note': cells[12],
            'inventoryPath': str(inventory_path),
        }
    return inventory, inventory_path


def _extract_calibration_history(path: Path) -> dict[str, str]:
    doc = Document(path)
    if len(doc.tables) < 2:
        return {}
    latest_cells: list[str] = []
    for row in doc.tables[1].rows[2:]:
        cells = [_clean_docx_text(cell.text) for cell in row.cells]
        if len(cells) < 9:
            continue
        if any(cells[1:]):
            latest_cells = cells
    if not latest_cells:
        return {}
    calib_method = ''
    if _cell_checked(latest_cells[3]):
        calib_method = '免校'
    elif _cell_checked(latest_cells[1]):
        calib_method = '內校'
    elif _cell_checked(latest_cells[2]):
        calib_method = '外校'
    note = latest_cells[8]
    return {
        'calibMethod': calib_method,
        'qualified': 'Y' if (_cell_checked(latest_cells[5]) and not _cell_checked(latest_cells[6])) else '',
        'checker': latest_cells[4],
        'nextCalibration': _normalize_iso_date(latest_cells[7]),
        'calibratedDate': _normalize_iso_date(note),
        'note': note,
    }


def _pick_latest_report(report_files: list[Path]) -> Path | None:
    if not report_files:
        return None
    return max(report_files, key=lambda item: item.stat().st_mtime)


def _latest_calibration_updates_by_instrument() -> dict[str, dict]:
    updates = {}
    try:
        records = ops_data.load_records('calibration')
    except Exception:
        return updates
    sorted_records = sorted(
        [item for item in records if str(item.get('instrumentId') or '').strip()],
        key=lambda item: (
            str(item.get('calibrationDate') or ''),
            str(item.get('updated_at') or ''),
            str(item.get('id') or ''),
        ),
    )
    for item in sorted_records:
        updates[str(item.get('instrumentId') or '').strip()] = item
    return updates


def load_burlan_calibration_instruments() -> dict:
    if not BURLAN_CALIBRATION_DIR.exists():
        return {
            'items': [],
            'source_path': str(BURLAN_CALIBRATION_DIR),
            'count': 0,
            'message': '找不到 9量測資源管理程序 資料夾',
            'mode': 'fallback',
            'latest_plan_path': '',
            'inventory_path': '',
            'manual_update_count': 0,
        }

    records_dir = BURLAN_CALIBRATION_DIR / '記錄'
    inventory_map, inventory_path = _load_calibration_inventory(records_dir)
    manual_updates = _latest_calibration_updates_by_instrument()
    record_files = sorted(
        path for path in records_dir.glob('9.2.*量規儀器履歷表*.docx')
        if '免校版' not in path.name
    )
    if not record_files and not inventory_map:
        return {
            'items': [],
            'source_path': str(BURLAN_CALIBRATION_DIR),
            'count': 0,
            'message': '找不到可用的量規儀器一覽表或履歷表',
            'mode': 'fallback',
            'latest_plan_path': '',
            'inventory_path': '',
            'manual_update_count': len(manual_updates),
        }

    annual_dirs = sorted(
        [path for path in records_dir.iterdir() if path.is_dir() and path.name.endswith('年度校正')],
        key=lambda item: item.name,
    )
    latest_annual_dir = annual_dirs[-1] if annual_dirs else None
    latest_plan_path = ''
    report_candidates = []
    if latest_annual_dir is not None:
        for file_path in latest_annual_dir.glob('9.3*.docx'):
            latest_plan_path = str(file_path)
            break
    for annual_dir in annual_dirs:
        report_candidates.extend(path for path in annual_dir.glob('9.4*.docx'))

    items = []
    handled_ids: set[str] = set()
    for path in record_files:
        pairs = _extract_docx_pairs(path)
        history = _extract_calibration_history(path)
        instrument_code = pairs.get('編號', '')
        inventory_item = inventory_map.get(instrument_code, {})
        instrument_name = pairs.get('儀器名稱', _extract_instrument_alias(path))
        frequency_label = inventory_item.get('frequencyLabel') or pairs.get('校驗頻率', '')
        alias = _extract_instrument_alias(path)
        related_reports = [
            report_path for report_path in report_candidates
            if (
                (instrument_code and f'({instrument_code})' in report_path.name)
                or (alias and f'({alias})' in report_path.name)
                or (alias and alias in report_path.name)
            )
        ]
        latest_report = _pick_latest_report(related_reports)
        calibrated_date = history.get('calibratedDate') or (datetime.fromtimestamp(latest_report.stat().st_mtime).date().isoformat() if latest_report else '')
        interval_days = _frequency_to_days(frequency_label)
        next_calibration = history.get('nextCalibration') or (add_days(calibrated_date, interval_days) if calibrated_date and interval_days else '')
        calib_method = inventory_item.get('calibMethod') or history.get('calibMethod') or ('外校' if latest_report else '')
        status = '免校正' if calib_method == '免校' else ('合格' if (latest_report or history.get('qualified')) else '待確認')
        instrument_id = instrument_code or inventory_item.get('id') or alias or path.stem
        manual_update = manual_updates.get(instrument_id)
        if manual_update:
            manual_calibration_date = str(manual_update.get('calibrationDate') or calibrated_date)
            manual_next_calibration = str(manual_update.get('nextCalibration') or '').strip()
            calibrated_date = manual_calibration_date
            next_calibration = manual_next_calibration or (add_days(manual_calibration_date, interval_days) if manual_calibration_date and interval_days else next_calibration)
            calib_method = str(manual_update.get('calibMethod') or calib_method)
            status = str(manual_update.get('status') or status)
        handled_ids.add(instrument_id)
        items.append(
            {
                'id': instrument_id,
                'name': instrument_name,
                'type': _guess_instrument_type(instrument_name),
                'location': pairs.get('使用單位', '') or inventory_item.get('location', ''),
                'keeper': pairs.get('保管人', '') or inventory_item.get('keeper', ''),
                'brand': pairs.get('廠商', '') or inventory_item.get('brand', ''),
                'model': inventory_item.get('model', '') or pairs.get('規格', ''),
                'serialNo': inventory_item.get('serialNo', ''),
                'calibMethod': calib_method,
                'calibratedDate': calibrated_date,
                'intervalDays': interval_days or 0,
                'status': status,
                'frequencyLabel': frequency_label,
                'registeredDate': _normalize_iso_date(pairs.get('登錄日期', '')),
                'nextCalibration': next_calibration,
                'needsMSA': False,
                'recordPath': str(path),
                'inventoryPath': inventory_item.get('inventoryPath', str(inventory_path) if inventory_path else ''),
                'latestReportPath': str(latest_report) if latest_report else '',
                'latestPlanPath': latest_plan_path,
                'sourceSystem': 'burlan_calibration_records',
                'manualRecordId': str(manual_update.get('id') or '') if manual_update else '',
                'manualUpdatedAt': str(manual_update.get('updated_at') or '') if manual_update else '',
                'manualNote': str(manual_update.get('note') or '') if manual_update else '',
                'manualOperator': str(manual_update.get('operator') or '') if manual_update else '',
            }
        )

    for instrument_code, inventory_item in inventory_map.items():
        if instrument_code in handled_ids:
            continue
        frequency_label = inventory_item.get('frequencyLabel', '')
        calib_method = inventory_item.get('calibMethod', '')
        manual_update = manual_updates.get(instrument_code)
        calibrated_date = str(manual_update.get('calibrationDate') or '') if manual_update else ''
        next_calibration = str(manual_update.get('nextCalibration') or '').strip() if manual_update else ''
        if manual_update:
            calib_method = str(manual_update.get('calibMethod') or calib_method)
            interval_days = _frequency_to_days(frequency_label) or 0
            if calibrated_date and interval_days:
                next_calibration = next_calibration or add_days(calibrated_date, interval_days)
        status = '免校正' if calib_method == '免校' else ('合格' if manual_update else '待確認')
        if manual_update:
            status = str(manual_update.get('status') or status)
        items.append(
            {
                'id': instrument_code,
                'name': inventory_item.get('name', instrument_code),
                'type': _guess_instrument_type(inventory_item.get('name', instrument_code)),
                'location': inventory_item.get('location', ''),
                'keeper': inventory_item.get('keeper', ''),
                'brand': inventory_item.get('brand', ''),
                'model': inventory_item.get('model', '') or inventory_item.get('spec', ''),
                'serialNo': inventory_item.get('serialNo', ''),
                'calibMethod': calib_method,
                'calibratedDate': calibrated_date,
                'intervalDays': _frequency_to_days(frequency_label) or 0,
                'status': status,
                'frequencyLabel': frequency_label,
                'registeredDate': '',
                'nextCalibration': next_calibration,
                'needsMSA': False,
                'recordPath': '',
                'inventoryPath': inventory_item.get('inventoryPath', str(inventory_path) if inventory_path else ''),
                'latestReportPath': '',
                'latestPlanPath': latest_plan_path,
                'sourceSystem': 'burlan_calibration_records',
                'manualRecordId': str(manual_update.get('id') or '') if manual_update else '',
                'manualUpdatedAt': str(manual_update.get('updated_at') or '') if manual_update else '',
                'manualNote': str(manual_update.get('note') or '') if manual_update else '',
                'manualOperator': str(manual_update.get('operator') or '') if manual_update else '',
            }
        )

    items.sort(key=lambda item: item['id'])
    return {
        'items': items,
        'source_path': str(records_dir),
        'count': len(items),
        'message': '已載入 9.1量規儀器一覽表與 9.2量規儀器履歷表',
        'mode': 'records',
        'latest_plan_path': latest_plan_path,
        'inventory_path': str(inventory_path) if inventory_path else '',
        'manual_update_count': len(manual_updates),
    }


def _pick_first_existing_file(path: Path, patterns: list[str]) -> Path | None:
    for pattern in patterns:
        for file_path in sorted(path.glob(pattern)):
            if file_path.is_file():
                return file_path
    return None


def _extract_month_number(value: str) -> int | None:
    match = re.search(r'(\d{1,2})\s*月', _clean_docx_text(value))
    if not match:
        return None
    month = int(match.group(1))
    return month if 1 <= month <= 12 else None


def _extract_month_day_to_iso(value: str, year: int) -> str:
    text = _clean_docx_text(value).replace(' ', '')
    match = re.search(r'(\d{1,2})/(\d{1,2})', text)
    if not match:
        return ''
    month = int(match.group(1))
    day = int(match.group(2))
    try:
        return datetime(year, month, day).date().isoformat()
    except ValueError:
        return ''


def _roc_year_text_to_ad(value: str) -> int | None:
    match = re.search(r'(\d{3,4})', _clean_docx_text(value))
    if not match:
        return None
    year = int(match.group(1))
    return year + 1911 if year < 1911 else year


def _first_value_after_label(cells: list[str], label: str) -> str:
    found = False
    for cell in cells:
        text = _clean_docx_text(cell)
        if not found:
            if text == label:
                found = True
            continue
        if text and text != label:
            return text
    return ''


def _unique_non_empty(values: list[str]) -> list[str]:
    results: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = _clean_docx_text(value)
        if not text or text in seen:
            continue
        seen.add(text)
        results.append(text)
    return results


def _equipment_manual_records_by_id() -> dict[str, dict]:
    try:
        records = ops_data.load_records('equipment')
    except Exception:
        return {}
    return {
        str(item.get('id') or '').strip(): item
        for item in records
        if str(item.get('id') or '').strip()
    }


def _supplier_manual_records_by_name() -> tuple[dict[str, dict], list[dict]]:
    try:
        records = ops_data.load_records('supplier')
    except Exception:
        return {}, []
    by_name: dict[str, dict] = {}
    extras: list[dict] = []
    for item in records:
        name = str(item.get('name') or '').strip()
        if name:
            by_name[name] = item
        else:
            extras.append(item)
    return by_name, records


def _is_legacy_equipment_seed(item: dict) -> bool:
    return str(item.get('id') or '').strip() in LEGACY_EQUIPMENT_IDS


def _is_legacy_supplier_seed(item: dict) -> bool:
    item_id = str(item.get('id') or '').strip()
    name = str(item.get('name') or '').strip()
    return item_id in LEGACY_SUPPLIER_IDS and name in LEGACY_SUPPLIER_NAMES


def _combine_equipment_history(source_history: list[dict], manual_history: list[dict]) -> list[dict]:
    merged: list[dict] = []
    seen: set[tuple[str, str, str]] = set()
    for item in list(manual_history) + list(source_history):
        key = (
            str(item.get('date') or '').strip(),
            str(item.get('operator') or '').strip(),
            '|'.join(str(part).strip() for part in (item.get('items') or [])),
        )
        if key in seen:
            continue
        seen.add(key)
        merged.append(item)
    merged.sort(key=lambda item: str(item.get('date') or ''), reverse=True)
    return merged


def _combine_supplier_history(source_history: list[dict], manual_history: list[dict]) -> list[dict]:
    merged: list[dict] = []
    seen: set[tuple[str, str, str]] = set()
    for item in list(manual_history) + list(source_history):
        key = (
            str(item.get('date') or '').strip(),
            str(item.get('score') or '').strip(),
            str(item.get('result') or '').strip(),
            str(item.get('operator') or '').strip(),
        )
        if key in seen:
            continue
        seen.add(key)
        merged.append(item)
    merged.sort(key=lambda item: str(item.get('date') or ''), reverse=True)
    return merged


def _load_equipment_inventory(records_dir: Path) -> tuple[dict[str, dict], Path | None]:
    inventory_path = _pick_first_existing_file(records_dir, ['8.1*.xlsx', '8.1*.xlsm'])
    if inventory_path is None:
        return {}, None
    workbook = load_workbook(inventory_path, data_only=True)
    try:
        worksheet = workbook[workbook.sheetnames[0]]
        inventory: dict[str, dict] = {}
        for row in worksheet.iter_rows(min_row=3, values_only=True):
            row_values = list(row[:8]) + [None] * (8 - len(row[:8]))
            equipment_id = _clean_docx_text(row_values[1])
            equipment_name = _clean_docx_text(row_values[2])
            if not equipment_id or not equipment_name:
                continue
            inventory[equipment_id] = {
                'id': equipment_id,
                'name': equipment_name,
                'model': _clean_docx_text(row_values[3]),
                'vendor': _clean_docx_text(row_values[4]),
                'quantity': _clean_docx_text(row_values[5]),
                'registeredDate': _normalize_iso_date(_clean_docx_text(row_values[6])),
                'location': '管理部',
                'inventoryPath': str(inventory_path),
            }
    finally:
        workbook.close()
    return inventory, inventory_path


def _extract_equipment_record_history(path: Path) -> dict:
    try:
        doc = Document(path)
    except Exception:
        return {}
    if not doc.tables:
        return {}
    table = doc.tables[0]
    rows = [[_clean_docx_text(cell.text) for cell in row.cells] for row in table.rows]
    if not rows:
        return {}
    header = rows[0]
    equipment_id = _first_value_after_label(header, '設備編號')
    name_match = re.search(r'\(([^)]+)\)', path.stem)
    equipment_name = _first_value_after_label(header, '設備名稱') or (name_match.group(1) if name_match else path.stem)
    owner = _first_value_after_label(header, '保管人')
    location = _first_value_after_label(header, '單位')

    next_items: list[str] = []
    maintenance_dates: list[str] = []
    year = _roc_year_text_to_ad(path.parent.name) or datetime.fromtimestamp(path.stat().st_mtime).year

    for row in rows[2:]:
        row_text = ''.join(row)
        if '保養時間' in row_text:
            for idx, cell in enumerate(row):
                date_value = _extract_month_day_to_iso(cell, year)
                if date_value:
                    maintenance_dates.append(date_value)
            continue
        if '備註' in row_text:
            continue
        item_label = ''
        for cell in row[1:]:
            text = _clean_docx_text(cell)
            if text and '每月保養' not in text:
                item_label = text
                break
        if item_label:
            next_items.append(item_label)

    history = [
        {
            'id': f"{equipment_id or equipment_name}-SRC-{index + 1:03d}",
            'date': maintenance_date,
            'operator': owner,
            'remark': '原始設備保養紀錄',
            'items': _unique_non_empty(next_items),
        }
        for index, maintenance_date in enumerate(sorted(set(maintenance_dates), reverse=True))
    ]
    return {
        'id': equipment_id,
        'name': equipment_name,
        'owner': owner,
        'location': location,
        'nextItems': _unique_non_empty(next_items),
        'maintenanceHistory': history,
        'lastMaintenance': history[0]['date'] if history else '',
        'recordPath': str(path),
    }


def load_burlan_equipment_records() -> dict:
    if not BURLAN_EQUIPMENT_DIR.exists():
        return {
            'items': [],
            'source_path': str(BURLAN_EQUIPMENT_DIR),
            'count': 0,
            'message': '找不到 8設施設備管理程序 資料夾',
            'mode': 'fallback',
        }

    records_dir = BURLAN_EQUIPMENT_DIR / '記錄'
    inventory_map, inventory_path = _load_equipment_inventory(records_dir)
    manual_map = _equipment_manual_records_by_id()
    annual_dirs = sorted([path for path in records_dir.iterdir() if path.is_dir() and '設備保養紀錄表' in path.name], key=lambda item: item.name)
    equipment_records: dict[str, dict] = {}
    for annual_dir in annual_dirs:
        for doc_path in sorted(annual_dir.iterdir()):
            if not doc_path.is_file() or doc_path.suffix.lower() != '.docx':
                continue
            record = _extract_equipment_record_history(doc_path)
            equipment_id = str(record.get('id') or '').strip()
            if not equipment_id:
                continue
            existing = equipment_records.get(equipment_id, {})
            combined_history = _combine_equipment_history(existing.get('maintenanceHistory', []), record.get('maintenanceHistory', []))
            equipment_records[equipment_id] = {
                **existing,
                **record,
                'maintenanceHistory': combined_history,
                'lastMaintenance': combined_history[0]['date'] if combined_history else record.get('lastMaintenance', ''),
                'nextItems': _unique_non_empty(list(existing.get('nextItems', [])) + list(record.get('nextItems', []))),
            }

    items: list[dict] = []
    handled_ids: set[str] = set()
    for equipment_id, inventory_item in sorted(inventory_map.items()):
        record_item = equipment_records.get(equipment_id, {})
        manual_item = manual_map.get(equipment_id, {})
        combined_history = _combine_equipment_history(record_item.get('maintenanceHistory', []), manual_item.get('maintenanceHistory', []))
        last_maintenance = str(manual_item.get('lastMaintenance') or record_item.get('lastMaintenance') or '')
        if combined_history and not last_maintenance:
            last_maintenance = combined_history[0].get('date', '')
        interval_days = int(manual_item.get('intervalDays') or 30)
        next_maintenance = str(manual_item.get('nextMaintenance') or (add_days(last_maintenance, interval_days) if last_maintenance else ''))
        items.append({
            'id': equipment_id,
            'name': str(manual_item.get('name') or record_item.get('name') or inventory_item.get('name') or equipment_id),
            'location': str(manual_item.get('location') or record_item.get('location') or inventory_item.get('location') or '管理部'),
            'owner': str(manual_item.get('owner') or record_item.get('owner') or ''),
            'intervalDays': interval_days,
            'lastMaintenance': last_maintenance,
            'nextMaintenance': next_maintenance,
            'vendor': str(manual_item.get('vendor') or inventory_item.get('vendor') or ''),
            'model': str(manual_item.get('model') or inventory_item.get('model') or ''),
            'status': str(manual_item.get('status') or ('待安排' if not last_maintenance else '正常')),
            'note': str(manual_item.get('note') or ''),
            'nextItems': _unique_non_empty(list(manual_item.get('nextItems') or []) + list(record_item.get('nextItems') or [])),
            'maintenanceHistory': combined_history,
            'source_file': str(record_item.get('recordPath') or inventory_item.get('inventoryPath') or ''),
            'sourceSystem': 'burlan_equipment_records',
            'inventoryPath': inventory_item.get('inventoryPath', str(inventory_path) if inventory_path else ''),
            'recordPath': str(record_item.get('recordPath') or ''),
        })
        handled_ids.add(equipment_id)

    for equipment_id, manual_item in manual_map.items():
        if equipment_id in handled_ids or _is_legacy_equipment_seed(manual_item):
            continue
        items.append({
            **manual_item,
            'sourceSystem': str(manual_item.get('sourceSystem') or 'manual_equipment_record'),
        })

    items.sort(key=lambda item: item.get('id', ''))
    return {
        'items': items,
        'source_path': str(records_dir),
        'count': len(items),
        'message': '已載入 8.1機器設備一覽表與設備保養紀錄',
        'mode': 'records',
        'inventory_path': str(inventory_path) if inventory_path else '',
        'manual_update_count': len([item for item in manual_map.values() if not _is_legacy_equipment_seed(item)]),
    }


def _supplier_score_to_result(score: int) -> str:
    if score >= 90:
        return '優良'
    if score >= 80:
        return '合格'
    if score >= 70:
        return '條件合格'
    return '不合格'


def _read_supplier_satisfaction_xls_rows(path: Path) -> list[list[str]]:
    try:
        import pythoncom
        import win32com.client as win32_client
    except Exception:
        return []

    temp_copy_path: Path | None = None
    workbook = None
    excel = None
    pythoncom.CoInitialize()
    try:
        suffix = path.suffix or '.xls'
        with tempfile.NamedTemporaryFile(prefix='burlan_supplier_', suffix=suffix, delete=False) as temp_file:
            temp_file.write(path.read_bytes())
            temp_copy_path = Path(temp_file.name)

        excel = win32_client.DispatchEx('Excel.Application')
        excel.Visible = False
        excel.DisplayAlerts = False
        workbook = excel.Workbooks.Open(str(temp_copy_path))
        sheet = workbook.Worksheets(1)
        used_range = sheet.UsedRange
        max_rows = int(used_range.Rows.Count or 0)
        max_cols = min(int(used_range.Columns.Count or 0), 8)
        rows: list[list[str]] = []
        for row_index in range(1, max_rows + 1):
            row_values = [_clean_docx_text(sheet.Cells(row_index, column_index).Text) for column_index in range(1, max_cols + 1)]
            rows.append(row_values)
        return rows
    except Exception:
        return []
    finally:
        if workbook is not None:
            try:
                workbook.Close(False)
            except Exception:
                pass
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        if temp_copy_path and temp_copy_path.exists():
            try:
                temp_copy_path.unlink()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _extract_supplier_satisfaction_summaries(path: Path) -> list[dict]:
    rows = _read_supplier_satisfaction_xls_rows(path)
    if len(rows) < 4:
        return []

    headers = rows[2]
    survey_year = _roc_year_text_to_ad(path.stem) or datetime.fromtimestamp(path.stat().st_mtime).year
    survey_date = f'{survey_year}-12-31'
    summaries: list[dict] = []
    for row in rows[3:]:
        supplier_name = _clean_docx_text(row[0] if len(row) > 0 else '')
        if not supplier_name:
            continue
        if supplier_name in {'客戶', '說明', '說  明'} or '製表' in supplier_name:
            continue
        average_text = _clean_docx_text(row[6] if len(row) > 6 else '')
        score_match = re.search(r'(\d+)', average_text)
        if not score_match:
            continue
        score = int(score_match.group(1))
        result = _clean_docx_text(row[7] if len(row) > 7 else '') or _supplier_score_to_result(score)
        issues: list[str] = []
        for header, rating in zip(headers[1:6], row[1:6]):
            normalized_rating = _clean_docx_text(rating)
            if not normalized_rating or normalized_rating in {'非常滿意', '滿意'}:
                continue
            issues.append(f'{_clean_docx_text(header)}：{normalized_rating}')
        summaries.append({
            'name': supplier_name,
            'date': survey_date,
            'score': score,
            'result': result,
            'issues': _unique_non_empty(issues),
            'source_file': str(path),
            'remark': f'來源：{path.stem}',
        })
    return summaries


def _append_supplier_grouped_entry(
    grouped: dict[str, dict],
    supplier_name: str,
    history_entry: dict,
    source_file: str,
    note: str = '',
) -> None:
    current = grouped.setdefault(
        supplier_name,
        {
            'name': supplier_name,
            'category': '採購供應商',
            'contact': '',
            'lastEvalDate': '',
            'evalScore': 0,
            'evalResult': '',
            'evalIntervalDays': 183,
            'issues': [],
            'note': '',
            'source_file': source_file,
            'evalHistory': [],
        },
    )
    current['evalHistory'].append(history_entry)
    if note and not current.get('note'):
        current['note'] = note
    current_date = str(current.get('lastEvalDate') or '')
    current_score = int(current.get('evalScore') or 0)
    if history_entry['date'] > current_date or (history_entry['date'] == current_date and int(history_entry.get('score') or 0) >= current_score):
        current['lastEvalDate'] = history_entry['date']
        current['evalScore'] = int(history_entry.get('score') or 0)
        current['evalResult'] = str(history_entry.get('result') or '')
        current['issues'] = list(history_entry.get('issues') or [])
        current['source_file'] = source_file


def _extract_latest_date_from_row_cells(cells: list[str]) -> str:
    dates = [_normalize_iso_date(cell) for cell in cells if _normalize_iso_date(cell)]
    return max(dates) if dates else ''


def _extract_supplier_doc_summary(path: Path) -> dict:
    try:
        doc = Document(path)
    except Exception:
        return {}
    if not doc.tables:
        return {}
    rows = [[_clean_docx_text(cell.text) for cell in row.cells] for row in doc.tables[0].rows]
    if not rows:
        return {}
    supplier_candidates = _unique_non_empty(rows[0][2:])
    supplier_name = supplier_candidates[0] if supplier_candidates else path.stem
    body_rows = rows[2:20]
    dated_rows = [row for row in body_rows if any(_normalize_iso_date(cell) for cell in row)]
    latest_date = ''
    for row in dated_rows:
        candidate = _extract_latest_date_from_row_cells(row)
        if candidate and candidate > latest_date:
            latest_date = candidate
    issues: list[str] = []
    for row in body_rows:
        if not row:
            continue
        issue_name = _clean_docx_text(row[1] if len(row) > 1 else '')
        markers = row[3:6] if len(row) >= 6 else row[2:]
        if issue_name and any(_cell_checked(marker) for marker in markers):
            issues.append(issue_name)
    issue_count = len(issues)
    score = max(0, 100 - issue_count * 20)
    result = _supplier_score_to_result(score)
    return {
        'name': supplier_name,
        'date': latest_date,
        'score': score,
        'result': result,
        'issues': _unique_non_empty(issues),
        'source_file': str(path),
        'remark': f'來源：{path.stem}',
    }


def load_burlan_supplier_records() -> dict:
    supplier_dir_exists = BURLAN_SUPPLIER_DIR.exists()
    third_level_dir_exists = BURLAN_THIRD_LEVEL_DIR.exists()
    if not supplier_dir_exists and not third_level_dir_exists:
        return {
            'items': [],
            'source_path': str(BURLAN_SUPPLIER_DIR),
            'count': 0,
            'message': '找不到 12採購及供應商管理程序 與 三階文件 資料夾',
            'mode': 'fallback',
        }

    records_dir = BURLAN_SUPPLIER_DIR / '記錄'
    doc_files = sorted(records_dir.glob('12.2*.docx')) if records_dir.exists() else []
    survey_files = sorted(BURLAN_THIRD_LEVEL_DIR.glob('供應商滿意度清單(*年度).xls')) if third_level_dir_exists else []
    if not doc_files and not survey_files:
        return {
            'items': [],
            'source_path': str(records_dir),
            'count': 0,
            'message': '找不到可用的供應商評鑑表或供應商滿意度清單',
            'mode': 'fallback',
        }

    grouped: dict[str, dict] = {}
    for path in survey_files:
        summaries = _extract_supplier_satisfaction_summaries(path)
        for index, summary in enumerate(summaries, 1):
            supplier_name = str(summary.get('name') or '').strip()
            if not supplier_name:
                continue
            history_entry = {
                'id': f"{path.stem}-SAT-{index:03d}",
                'date': str(summary.get('date') or ''),
                'score': int(summary.get('score') or 0),
                'result': str(summary.get('result') or ''),
                'operator': '管理部',
                'remark': str(summary.get('remark') or ''),
                'issues': list(summary.get('issues') or []),
            }
            _append_supplier_grouped_entry(grouped, supplier_name, history_entry, str(path), note='來源：年度供應商滿意度清單')

    for path in doc_files:
        summary = _extract_supplier_doc_summary(path)
        supplier_name = str(summary.get('name') or '').strip()
        if not supplier_name:
            continue
        history_entry = {
            'id': f"{path.stem}-SRC",
            'date': str(summary.get('date') or ''),
            'score': int(summary.get('score') or 0),
            'result': str(summary.get('result') or ''),
            'operator': '管理部',
            'remark': str(summary.get('remark') or ''),
            'issues': list(summary.get('issues') or []),
        }
        _append_supplier_grouped_entry(grouped, supplier_name, history_entry, str(path))

    manual_by_name, manual_records = _supplier_manual_records_by_name()
    items: list[dict] = []
    handled_names: set[str] = set()
    for index, supplier_name in enumerate(sorted(grouped.keys()), 1):
        source_item = grouped[supplier_name]
        source_item['evalHistory'] = _combine_supplier_history(source_item.get('evalHistory', []), [])
        latest_source_history = source_item['evalHistory'][0] if source_item['evalHistory'] else {}
        if latest_source_history:
            source_item['lastEvalDate'] = str(latest_source_history.get('date') or source_item.get('lastEvalDate') or '')
            source_item['evalScore'] = int(latest_source_history.get('score') or source_item.get('evalScore') or 0)
            source_item['evalResult'] = str(latest_source_history.get('result') or source_item.get('evalResult') or '')
            source_item['issues'] = list(latest_source_history.get('issues') or source_item.get('issues') or [])
        manual_item = manual_by_name.get(supplier_name, {})
        combined_history = _combine_supplier_history(source_item.get('evalHistory', []), manual_item.get('evalHistory', []))
        last_eval_date = str(manual_item.get('lastEvalDate') or source_item.get('lastEvalDate') or '')
        latest_history = combined_history[0] if combined_history else {}
        eval_score = int(manual_item.get('evalScore') or source_item.get('evalScore') or latest_history.get('score') or 0)
        eval_result = str(manual_item.get('evalResult') or source_item.get('evalResult') or latest_history.get('result') or '')
        issues = _unique_non_empty(list(manual_item.get('issues') or []) + list(source_item.get('issues') or []))
        items.append({
            'id': str(manual_item.get('id') or f'SUP-BR-{index:03d}'),
            'name': supplier_name,
            'category': str(manual_item.get('category') or source_item.get('category') or '採購供應商'),
            'contact': str(manual_item.get('contact') or ''),
            'lastEvalDate': last_eval_date,
            'evalScore': eval_score,
            'evalResult': eval_result or _supplier_score_to_result(eval_score),
            'evalIntervalDays': int(manual_item.get('evalIntervalDays') or source_item.get('evalIntervalDays') or 183),
            'issues': issues,
            'note': str(manual_item.get('note') or source_item.get('note') or ''),
            'source_file': str(source_item.get('source_file') or ''),
            'evalHistory': combined_history,
            'sourceSystem': 'burlan_supplier_records',
        })
        handled_names.add(supplier_name)

    for manual_item in manual_records:
        supplier_name = str(manual_item.get('name') or '').strip()
        if supplier_name in handled_names or _is_legacy_supplier_seed(manual_item):
            continue
        items.append({
            **manual_item,
            'sourceSystem': str(manual_item.get('sourceSystem') or 'manual_supplier_record'),
        })

    items.sort(key=lambda item: item.get('name', ''))
    return {
        'items': items,
        'source_path': str(records_dir),
        'count': len(items),
        'message': '已載入 年度供應商滿意度清單 與 12.2供應商評鑑表',
        'mode': 'records',
        'survey_source_path': str(BURLAN_THIRD_LEVEL_DIR),
        'manual_update_count': len([item for item in manual_records if not _is_legacy_supplier_seed(item)]),
    }


def save_ops_records(kind: str):
    body = json_body()
    records = body.get('records')
    replace_source_file = str(body.get('replace_source_file') or '').strip()
    if records is None:
        record = body.get('record')
        records = record if isinstance(record, list) else [record or body]
    if not isinstance(records, list) or not records:
        return json_error('No records provided.')
    items, saved = ops_data.upsert_records(
        kind,
        [record for record in records if isinstance(record, dict)],
        replace_source_file=replace_source_file,
    )
    payload = {'items': items, 'saved': saved}
    if kind == 'environment':
        payload['summary'] = ops_data.summarize_environment(items)
    return jsonify(payload)


def _is_closed_nonconformance(status: str) -> bool:
    return str(status or '').strip() in {'已關閉', '已結案', 'Closed'}


def enrich_audit_plans_with_nonconformance_summary(items=None):
    audit_plans = [dict(item) for item in (items if items is not None else load_burlan_audit_plans().get('items', []))]
    nonconformances = ops_data.load_records('nonconformance')
    grouped = {}
    for item in nonconformances:
        audit_id = str(item.get('sourceAuditPlanId') or '').strip()
        if not audit_id:
            continue
        grouped.setdefault(audit_id, []).append(item)

    enriched = []
    for plan in audit_plans:
        linked_items = grouped.get(str(plan.get('id') or '').strip(), [])
        open_items = [item for item in linked_items if not _is_closed_nonconformance(item.get('status'))]
        closed_items = [item for item in linked_items if _is_closed_nonconformance(item.get('status'))]
        overdue_items = [
            item for item in open_items
            if str(item.get('dueDate') or '').strip() and str(item.get('dueDate')) < datetime.now().date().isoformat()
        ]
        close_dates = sorted(str(item.get('closeDate') or '').strip() for item in closed_items if str(item.get('closeDate') or '').strip())
        enriched_plan = dict(plan)
        enriched_plan.update({
            'linkedNcCount': len(linked_items),
            'openLinkedNcCount': len(open_items),
            'closedLinkedNcCount': len(closed_items),
            'overdueLinkedNcCount': len(overdue_items),
            'latestLinkedNcCloseDate': close_dates[-1] if close_dates else '',
        })
        enriched.append(enriched_plan)
    return enriched


def serve_managed_file(stored_path: str, as_attachment: bool):
    file_path = ops_data.get_serving_path(stored_path)
    if file_path is None:
        abort(404)
    return send_file(file_path, as_attachment=as_attachment, download_name=file_path.name)


def google_status_payload() -> dict:
    return server_handlers_integrations.google_status_payload(
        load_google_config,
        load_google_tokens,
        google_configured,
        google_redirect_uri,
    )


def build_event_payload(item: dict) -> dict:
    return server_handlers_integrations.build_event_payload(item, add_days)


def add_days(date_str: str, days: int) -> str:
    from datetime import date, timedelta

    base = date.fromisoformat(date_str)
    return (base + timedelta(days=days)).isoformat()


def resolve_public_static_file(filename: str) -> Path | None:
    cleaned = filename.strip().lstrip('/').replace('\\', '/')
    if not cleaned or cleaned.startswith('.'):
        return None
    if cleaned in PUBLIC_STATIC_FILES:
        path = PUBLIC_STATIC_FILES[cleaned]
        return path if path.exists() and path.is_file() else None

    for prefix, root in PUBLIC_STATIC_DIRS.items():
        if cleaned == prefix or not cleaned.startswith(prefix + '/'):
            continue
        relative = cleaned[len(prefix) + 1 :]
        candidate = (root / relative).resolve()
        try:
            candidate.relative_to(root.resolve())
        except ValueError:
            return None
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def startup_security_warnings(host: str) -> list[str]:
    warnings = []
    public_private_files = public_root_contains_private_files()
    if public_private_files:
        warnings.append('Private config files are still present in the public project root.')
    if host == '0.0.0.0':
        warnings.append('Server is bound to 0.0.0.0. Other machines on the network may reach this site.')
    if not os.environ.get('FLASK_SECRET_KEY'):
        warnings.append('FLASK_SECRET_KEY is not provided. Using a generated local secret file instead.')
    return warnings


@app.route('/')
def index():
    return send_from_directory(BASE_DIR, 'index.html')


@app.route('/<path:filename>')
def static_files(filename):
    filepath = resolve_public_static_file(filename)
    if filepath is not None:
        return send_from_directory(filepath.parent, filepath.name)
    abort(404)


@app.route('/api/files/view')
def api_view_file():
    return server_handlers_documents.api_view_file(serve_managed_file)


@app.route('/api/files/download')
def api_download_file():
    return server_handlers_documents.api_download_file(serve_managed_file)


@app.route('/api/files/preview-text')
def api_preview_text_file():
    return server_handlers_documents.api_preview_text_file()


@app.route('/api/burlan/master-documents', methods=['GET'])
def api_burlan_master_documents():
    try:
        return server_handlers_documents.api_burlan_master_documents(load_burlan_master_documents)
    except Exception as exc:
        traceback.print_exc()
        return json_error(str(exc), 500)


@app.route('/api/burlan/quality-objectives', methods=['GET'])
def api_burlan_quality_objectives():
    try:
        return server_handlers_documents.api_burlan_quality_objectives(load_burlan_quality_objectives)
    except Exception as exc:
        traceback.print_exc()
        return json_error(str(exc), 500)


@app.route('/api/burlan/calibration-instruments', methods=['GET'])
def api_burlan_calibration_instruments():
    try:
        return server_handlers_documents.api_burlan_calibration_instruments(load_burlan_calibration_instruments)
    except Exception as exc:
        traceback.print_exc()
        return json_error(str(exc), 500)


@app.route('/api/calibration-records', methods=['GET'])
def api_calibration_records_list():
    return server_handlers_operations.api_calibration_records_list()


@app.route('/api/calibration-records', methods=['POST'])
def api_calibration_records_save():
    return server_handlers_operations.api_calibration_records_save(
        json_body,
        json_error,
        load_burlan_calibration_instruments,
    )


@app.route('/api/training-records', methods=['GET'])
def api_training_records_list():
    return server_handlers_operations.api_training_records_list()


@app.route('/api/training-records', methods=['POST'])
def api_training_records_save():
    return server_handlers_operations.api_training_records_save(save_ops_records, json_error)


@app.route('/api/training-records/<record_id>', methods=['DELETE'])
def api_training_records_delete(record_id):
    return server_handlers_operations.api_training_records_delete(record_id, json_error)


@app.route('/api/training-records/<record_id>/delete', methods=['POST'])
def api_training_records_delete_post(record_id):
    return api_training_records_delete(record_id)


@app.route('/api/equipment-records', methods=['GET'])
def api_equipment_records_list():
    return server_handlers_operations.api_equipment_records_list(load_burlan_equipment_records)


@app.route('/api/equipment-records', methods=['POST'])
def api_equipment_records_save():
    return server_handlers_operations.api_equipment_records_save(
        json_body,
        json_error,
        load_burlan_equipment_records,
    )


@app.route('/api/equipment-records/<record_id>', methods=['DELETE'])
def api_equipment_records_delete(record_id):
    return server_handlers_operations.api_equipment_records_delete(
        record_id,
        json_error,
        load_burlan_equipment_records,
    )


@app.route('/api/equipment-records/<record_id>/delete', methods=['POST'])
def api_equipment_records_delete_post(record_id):
    return api_equipment_records_delete(record_id)


@app.route('/api/supplier-records', methods=['GET'])
def api_supplier_records_list():
    return server_handlers_operations.api_supplier_records_list(load_burlan_supplier_records)


@app.route('/api/supplier-records', methods=['POST'])
def api_supplier_records_save():
    return server_handlers_operations.api_supplier_records_save(
        json_body,
        json_error,
        load_burlan_supplier_records,
    )


@app.route('/api/supplier-records/<record_id>', methods=['DELETE'])
def api_supplier_records_delete(record_id):
    return server_handlers_operations.api_supplier_records_delete(
        record_id,
        json_error,
        load_burlan_supplier_records,
    )


@app.route('/api/supplier-records/<record_id>/delete', methods=['POST'])
def api_supplier_records_delete_post(record_id):
    return api_supplier_records_delete(record_id)


@app.route('/api/nonconformances', methods=['GET'])
def api_nonconformances_list():
    return server_handlers_operations.api_nonconformances_list(
        enrich_audit_plans_with_nonconformance_summary,
    )


@app.route('/api/nonconformances', methods=['POST'])
def api_nonconformances_save():
    return server_handlers_operations.api_nonconformances_save(
        json_body,
        json_error,
        enrich_audit_plans_with_nonconformance_summary,
    )


@app.route('/api/nonconformances/<record_id>', methods=['DELETE'])
def api_nonconformances_delete(record_id):
    return server_handlers_operations.api_nonconformances_delete(
        record_id,
        json_error,
        enrich_audit_plans_with_nonconformance_summary,
    )


@app.route('/api/nonconformances/import', methods=['POST'])
def api_nonconformances_import():
    return server_handlers_operations.api_nonconformances_import(json_error)


@app.route('/api/audit-plans', methods=['GET'])
def api_audit_plans_list():
    return server_handlers_operations.api_audit_plans_list(
        load_burlan_audit_plans,
        enrich_audit_plans_with_nonconformance_summary,
    )


@app.route('/api/audit-plans', methods=['POST'])
def api_audit_plans_save():
    return server_handlers_operations.api_audit_plans_save(
        json_body,
        json_error,
        load_burlan_audit_plans,
        enrich_audit_plans_with_nonconformance_summary,
    )


@app.route('/api/audit-plans/<record_id>', methods=['DELETE'])
def api_audit_plans_delete(record_id):
    return server_handlers_operations.api_audit_plans_delete(
        record_id,
        json_error,
        load_burlan_audit_plans,
        enrich_audit_plans_with_nonconformance_summary,
    )


@app.route('/api/audit-plans/import', methods=['POST'])
def api_audit_plans_import():
    return server_handlers_operations.api_audit_plans_import(json_error)


@app.route('/api/audit-plans/<record_id>/attachments', methods=['GET'])
def api_audit_plan_attachments(record_id):
    return server_handlers_operations.api_audit_plan_attachments(
        record_id,
        json_error,
        load_burlan_audit_plans,
    )


@app.route('/api/audit-plans/year-bundle', methods=['GET'])
def api_audit_plan_year_bundle():
    return server_handlers_operations.api_audit_plan_year_bundle(
        json_error,
        build_audit_plan_year_bundle,
    )


@app.route('/api/environment-records', methods=['GET'])
def api_environment_records_list():
    return server_handlers_operations.api_environment_records_list()


@app.route('/api/environment-records', methods=['POST'])
def api_environment_records_save():
    return server_handlers_operations.api_environment_records_save(save_ops_records, json_error)


@app.route('/api/environment-records/<record_id>', methods=['DELETE'])
def api_environment_records_delete(record_id):
    return server_handlers_operations.api_environment_records_delete(record_id, json_error)


@app.route('/api/environment-records/import', methods=['POST'])
def api_environment_records_import():
    return server_handlers_operations.api_environment_records_import(json_error)


@app.route('/api/environment-records/delete-range', methods=['POST'])
def api_environment_records_delete_range():
    return server_handlers_operations.api_environment_records_delete_range(json_body, json_error)


@app.route('/api/production-records/read-existing', methods=['GET'])
def api_production_records_read_existing():
    return server_handlers_operations.api_production_records_read_existing(json_error)


@app.route('/api/production-records/import', methods=['POST'])
def api_production_records_import():
    return server_handlers_operations.api_production_records_import(json_error)


@app.route('/api/quality-records/read-existing', methods=['GET'])
def api_quality_records_read_existing():
    return server_handlers_operations.api_quality_records_read_existing(json_error)


@app.route('/api/quality-records/import', methods=['POST'])
def api_quality_records_import():
    return server_handlers_operations.api_quality_records_import(json_error)


@app.route('/api/generate', methods=['POST'])
def api_generate():
    return server_handlers_operations.api_generate(json_error)


@app.route('/api/record-engine/catalog', methods=['GET'])
def api_record_engine_catalog():
    return server_handlers_operations.api_record_engine_catalog()


@app.route('/api/record-engine/suggest', methods=['POST'])
def api_record_engine_suggest():
    return server_handlers_operations.api_record_engine_suggest()


@app.route('/api/record-engine/precheck', methods=['POST'])
def api_record_engine_precheck():
    return server_handlers_operations.api_record_engine_precheck()


@app.route('/api/record-engine/generate', methods=['POST'])
def api_record_engine_generate():
    return server_handlers_operations.api_record_engine_generate()


@app.route('/api/notion', methods=['POST'])
def api_notion():
    return server_handlers_integrations.api_notion()


@app.route('/api/shipment-draft/catalog', methods=['GET'])
def shipment_draft_catalog():
    return server_handlers_operations.shipment_draft_catalog()


@app.route('/api/shipment-draft/generate', methods=['POST'])
def shipment_draft_generate():
    return server_handlers_operations.shipment_draft_generate()


@app.route('/api/google-calendar/status', methods=['GET'])
def google_calendar_status():
    return server_handlers_integrations.google_calendar_status(google_status_payload)


@app.route('/api/google-calendar/config', methods=['POST'])
def google_calendar_config():
    return server_handlers_integrations.google_calendar_config(
        save_google_config,
        clear_google_config,
        clear_google_tokens,
        google_status_payload,
    )


@app.route('/api/google-calendar/auth/start', methods=['GET'])
def google_calendar_auth_start():
    return server_handlers_integrations.google_calendar_auth_start(
        load_google_config,
        google_configured,
        google_redirect_uri,
        GOOGLE_AUTH_URL,
        GOOGLE_SCOPES,
        secrets.token_urlsafe,
    )


@app.route('/api/google-calendar/oauth/callback', methods=['GET'])
def google_calendar_oauth_callback():
    return server_handlers_integrations.google_calendar_oauth_callback(
        load_google_config,
        exchange_google_code,
        save_google_tokens,
    )


@app.route('/api/google-calendar/events', methods=['POST'])
def google_calendar_events():
    return server_handlers_integrations.google_calendar_events(
        require_google_access_token,
        google_error_text,
        build_event_payload,
        GOOGLE_CALENDAR_EVENTS_URL,
    )


@app.route('/api/google-calendar/logout', methods=['POST'])
def google_calendar_logout():
    return server_handlers_integrations.google_calendar_logout(
        clear_google_tokens,
        google_status_payload,
    )


def kill_port(port: int) -> None:
    """啟動前清除佔用指定 port 的舊 Python 程序（Windows）"""
    import subprocess
    current_pid = os.getpid()
    try:
        result = subprocess.run(
            ['netstat', '-ano'], capture_output=True, text=True, timeout=5
        )
        for line in result.stdout.splitlines():
            if f':{port}' in line and 'LISTENING' in line:
                parts = line.split()
                if not parts:
                    continue
                try:
                    pid = int(parts[-1])
                except ValueError:
                    continue
                if pid != current_pid:
                    subprocess.run(
                        ['taskkill', '/F', '/PID', str(pid)],
                        capture_output=True, timeout=5
                    )
                    print(f'[cleanup] 已關閉舊程序 PID {pid}（port {port}）')
    except Exception:
        pass


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8888))
    host = os.environ.get('HOST', '127.0.0.1')
    kill_port(port)
    print(f'[server] http://{host}:{port}')
    print(f'[base] {BASE_DIR}')
    for warning in startup_security_warnings(host):
        print(f'[security-warning] {warning}')
    app.run(host=host, port=port, debug=False, use_reloader=False)
