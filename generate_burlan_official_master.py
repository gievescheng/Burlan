from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


BASE_DIR = Path(__file__).resolve().parent
DEFAULT_SOURCE_DIR = BASE_DIR.parent / "ISO 9001 SOP-柏連正式版"
DEFAULT_MAIN_LIST_PATH = (
    DEFAULT_SOURCE_DIR
    / "1 文件化資訊管制程序"
    / "記錄"
    / "1.2.1文件管制總覽表（程序文件2025年驗證版).docx"
)
DEFAULT_OUTPUT_PATH = DEFAULT_SOURCE_DIR / "柏連正式文件主清單.xlsx"

VALID_EXTENSIONS = {".pdf", ".docx", ".doc"}
EXTENSION_PRIORITY = {".pdf": 0, ".docx": 1, ".doc": 2}
REVIEW_OK = "主清單一致"
REVIEW_MANUAL = "待人工確認"


@dataclass
class MainListRow:
    序號: str
    文件編號: str
    文件名稱: str
    負責單位: str
    主清單版次: str
    主清單發行日期: str


@dataclass
class FileCandidate:
    path: Path
    version: str | None
    is_suspicious: bool

    @property
    def name(self) -> str:
        return self.path.name

    @property
    def suffix(self) -> str:
        return self.path.suffix.lower()


def normalize_version(text: str | None) -> str:
    if text is None:
        return ""
    cleaned = str(text).strip()
    if not cleaned:
        return ""
    try:
        return f"{float(cleaned):.1f}"
    except ValueError:
        return cleaned


def parse_version_from_name(stem: str) -> str | None:
    patterns = [
        r"\((\d+(?:\.\d+)?)\)",
        r"(\d+(?:\.\d+)?)版本",
        r"(?:程序|手冊)(\d+(?:\.\d+)?)$",
        r"(?<!\d)(\d+(?:\.\d+)?)(?!\d)",
    ]
    for pattern in patterns:
        matches = re.findall(pattern, stem)
        if not matches:
            continue
        for match in matches:
            try:
                value = float(match)
            except ValueError:
                continue
            if 0 < value < 20:
                return f"{value:.1f}"
    return None


def effective_version(file: FileCandidate, main_version: str) -> str:
    parsed = normalize_version(file.version)
    if parsed:
        return parsed
    if not file.is_suspicious and main_version == "1.0":
        return "1.0"
    return ""


def is_suspicious_name(name: str) -> bool:
    lowered = name.lower()
    keywords = [
        "conflicted copy",
        "~$",
        "自動儲存",
        "復原",
        "backup",
        "tmp",
    ]
    return any(keyword in lowered for keyword in keywords)


def load_main_list(path: Path) -> list[MainListRow]:
    doc = Document(str(path))
    if not doc.tables:
        raise ValueError(f"正式文件主清單來源沒有表格：{path}")

    rows: list[MainListRow] = []
    for row in doc.tables[0].rows[3:]:
        values = [cell.text.strip() for cell in row.cells]
        if len(values) < 6 or not values[0].isdigit():
            continue
        rows.append(
            MainListRow(
                序號=values[0],
                文件編號=values[1],
                文件名稱=values[2],
                負責單位=values[3],
                主清單版次=normalize_version(values[4]),
                主清單發行日期=values[5],
            )
        )
    return rows


def find_folder_for_code(source_dir: Path, code: str) -> Path:
    if code == "MM-01":
        folder_number = 0
    else:
        folder_number = int(code.split("-")[1])
    pattern = re.compile(rf"^{folder_number}(?:\s|[^0-9])")
    for child in sorted(source_dir.iterdir(), key=lambda item: item.name):
        if child.is_dir() and pattern.match(child.name):
            return child
    raise FileNotFoundError(f"找不到 {code} 對應的資料夾")


def collect_candidates(folder: Path) -> list[FileCandidate]:
    candidates: list[FileCandidate] = []
    for item in sorted(folder.iterdir(), key=lambda path: path.name.lower()):
        if not item.is_file():
            continue
        if item.suffix.lower() not in VALID_EXTENSIONS:
            continue
        candidates.append(
            FileCandidate(
                path=item,
                version=parse_version_from_name(item.stem),
                is_suspicious=is_suspicious_name(item.name),
            )
        )
    return candidates


def format_file_list(files: list[FileCandidate], suffixes: tuple[str, ...]) -> str:
    names = [file.name for file in files if file.suffix in suffixes and not file.is_suspicious]
    return "\n".join(names)


def choose_temp_file(files: list[FileCandidate], main_version: str) -> FileCandidate | None:
    exact = [
        file
        for file in files
        if not file.is_suspicious and effective_version(file, main_version) == main_version
    ]
    if not exact:
        return None
    return sorted(exact, key=lambda file: (EXTENSION_PRIORITY.get(file.suffix, 99), file.name.lower()))[0]


def summarize_review(files: list[FileCandidate], main_version: str, temp_file: FileCandidate | None) -> tuple[str, str]:
    reasons: list[str] = []
    numeric_versions = []
    for file in files:
        if file.is_suspicious:
            continue
        version = effective_version(file, main_version)
        if not version:
            continue
        try:
            numeric_versions.append(float(version))
        except ValueError:
            continue

    try:
        main_numeric = float(main_version)
    except ValueError:
        main_numeric = None

    if numeric_versions and main_numeric is not None:
        highest = max(numeric_versions)
        if highest > main_numeric:
            reasons.append(f"主清單為 {main_version}，但資料夾另有較新版 {highest:.1f}")

    suspicious_names = [file.name for file in files if file.is_suspicious]
    if suspicious_names:
        reasons.append("有需排除的暫存或衝突檔")

    if temp_file is None:
        reasons.append("找不到與主清單版次相符的安全候選檔")
    else:
        exact_pdf = any(
            file.suffix == ".pdf"
            and not file.is_suspicious
            and effective_version(file, main_version) == main_version
            for file in files
        )
        if not exact_pdf:
            reasons.append("目前找不到同版次 PDF")

    if reasons:
        return REVIEW_MANUAL, "；".join(reasons)

    return REVIEW_OK, "已找到與主清單版次一致的候選檔案"


def build_master_rows(source_dir: Path, main_list_path: Path) -> list[dict]:
    rows = load_main_list(main_list_path)
    results: list[dict] = []

    for row in rows:
        folder = find_folder_for_code(source_dir, row.文件編號)
        candidates = collect_candidates(folder)
        temp_file = choose_temp_file(candidates, row.主清單版次)
        review_status, review_reason = summarize_review(candidates, row.主清單版次, temp_file)

        results.append(
            {
                "文件編號": row.文件編號,
                "文件名稱": row.文件名稱,
                "文件類別": "品質手冊" if row.文件編號.startswith("MM-") else "程序文件",
                "負責單位": row.負責單位,
                "主清單版次": row.主清單版次,
                "主清單發行日期": row.主清單發行日期,
                "實際資料夾位置": str(folder),
                "找到的 PDF 檔": format_file_list(candidates, (".pdf",)),
                "找到的 Word 檔": format_file_list(candidates, (".docx", ".doc")),
                "暫定正式檔案": temp_file.name if temp_file else "",
                "判定狀態": review_status,
                "判定原因": review_reason,
                "是否有表單": "是" if (folder / "表單").is_dir() else "否",
                "是否有記錄": "是" if (folder / "記錄").is_dir() else "否",
                "是否納入系統": "是",
            }
        )

    return results


def apply_table_style(ws) -> None:
    title_fill = PatternFill("solid", fgColor="0F766E")
    manual_fill = PatternFill("solid", fgColor="FFF7ED")
    ok_fill = PatternFill("solid", fgColor="ECFDF5")
    header_fill = PatternFill("solid", fgColor="E2E8F0")
    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for cell in ws[1]:
        cell.fill = title_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        status = row[10].value
        fill = manual_fill if status == REVIEW_MANUAL else ok_fill
        for cell in row:
            cell.fill = fill
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    for header_cell in ws[1]:
        header_cell.fill = header_fill

    for header_cell in ws[1]:
        header_cell.fill = title_fill
        header_cell.font = Font(color="FFFFFF", bold=True)

    widths = {
        "A": 12,
        "B": 26,
        "C": 14,
        "D": 12,
        "E": 12,
        "F": 14,
        "G": 40,
        "H": 34,
        "I": 34,
        "J": 24,
        "K": 14,
        "L": 42,
        "M": 10,
        "N": 10,
        "O": 10,
    }
    for column, width in widths.items():
        ws.column_dimensions[column].width = width

    ws.row_dimensions[1].height = 24
    for index in range(2, ws.max_row + 1):
        ws.row_dimensions[index].height = 48

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(ws.max_column)}{ws.max_row}"


def build_summary_sheet(wb: Workbook, rows: list[dict], source_dir: Path, main_list_path: Path) -> None:
    ws = wb.create_sheet("整理摘要")
    ws["A1"] = "柏連正式文件主清單整理摘要"
    ws["A1"].font = Font(size=16, bold=True)
    ws["A3"] = "資料來源"
    ws["B3"] = str(main_list_path)
    ws["A4"] = "掃描資料夾"
    ws["B4"] = str(source_dir)
    ws["A5"] = "產生時間"
    ws["B5"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    ws["A7"] = "統計項目"
    ws["B7"] = "結果"
    ws["A8"] = "主清單總筆數"
    ws["B8"] = "=COUNTA(正式文件主清單!A2:A200)"
    ws["A9"] = "主清單一致"
    ws["B9"] = '=COUNTIF(正式文件主清單!K:K,"主清單一致")'
    ws["A10"] = "待人工確認"
    ws["B10"] = '=COUNTIF(正式文件主清單!K:K,"待人工確認")'
    ws["A11"] = "有表單"
    ws["B11"] = '=COUNTIF(正式文件主清單!M:M,"是")'
    ws["A12"] = "有記錄"
    ws["B12"] = '=COUNTIF(正式文件主清單!N:N,"是")'

    ws["A14"] = "一定要先人工確認的文件"
    ws["A15"] = "文件編號"
    ws["B15"] = "文件名稱"
    ws["C15"] = "判定原因"

    review_targets = {"MM-01", "MP-02", "MP-07"}
    target_rows = [row for row in rows if row["文件編號"] in review_targets]
    current_row = 16
    for row in target_rows:
        ws[f"A{current_row}"] = row["文件編號"]
        ws[f"B{current_row}"] = row["文件名稱"]
        ws[f"C{current_row}"] = row["判定原因"]
        current_row += 1

    ws["A21"] = "使用說明"
    ws["A22"] = "1. 主清單一致：代表主清單版次和資料夾候選檔一致，但仍建議人工覆核後再正式啟用。"
    ws["A23"] = "2. 待人工確認：代表有新版、衝突檔、缺少同版 PDF，或其他需要人判斷的情況。"
    ws["A24"] = "3. 暫定正式檔案：只是一個建議先查看的檔案，不代表已經正式定版。"

    fill_title = PatternFill("solid", fgColor="0F766E")
    fill_header = PatternFill("solid", fgColor="E2E8F0")
    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for cell in ws["A7:C7"][0]:
        cell.fill = fill_header
        cell.font = Font(bold=True)
        cell.border = border
        cell.alignment = Alignment(horizontal="center")

    for cell in ws["A15:C15"][0]:
        cell.fill = fill_header
        cell.font = Font(bold=True)
        cell.border = border
        cell.alignment = Alignment(horizontal="center")

    for cell in ("A1",):
        ws[cell].fill = fill_title
        ws[cell].font = Font(color="FFFFFF", size=16, bold=True)

    for row in ws.iter_rows(min_row=7, max_row=max(24, current_row - 1), min_col=1, max_col=3):
        for cell in row:
            if cell.value is None:
                continue
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 70


def build_instruction_sheet(wb: Workbook) -> None:
    ws = wb.create_sheet("欄位說明")
    rows = [
        ("欄位", "意思"),
        ("文件編號", "例如 MM-01、MP-07，用來分辨每一份正式程序文件。"),
        ("主清單版次", "正式清單上目前寫的版次。"),
        ("找到的 PDF 檔 / Word 檔", "系統在對應資料夾內找到的候選檔案。"),
        ("暫定正式檔案", "目前建議先看的檔案，不代表已經完成正式定版。"),
        ("判定狀態", "用來快速看這份文件現在是不是需要人工判斷。"),
        ("判定原因", "說明為什麼需要人工確認，或為什麼目前看起來一致。"),
        ("是否納入系統", "這份文件是否屬於柏連 QMS 正式範圍。"),
    ]
    for row in rows:
        ws.append(row)

    fill_header = PatternFill("solid", fgColor="0F766E")
    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for cell in ws[1]:
        cell.fill = fill_header
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center")
        cell.border = border

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 88


def write_workbook(rows: list[dict], output_path: Path, source_dir: Path, main_list_path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "正式文件主清單"

    headers = [
        "文件編號",
        "文件名稱",
        "文件類別",
        "負責單位",
        "主清單版次",
        "主清單發行日期",
        "實際資料夾位置",
        "找到的 PDF 檔",
        "找到的 Word 檔",
        "暫定正式檔案",
        "判定狀態",
        "判定原因",
        "是否有表單",
        "是否有記錄",
        "是否納入系統",
    ]
    ws.append(headers)
    for row in rows:
        ws.append([row[header] for header in headers])

    apply_table_style(ws)
    build_summary_sheet(wb, rows, source_dir, main_list_path)
    build_instruction_sheet(wb)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="產生柏連正式文件主清單 Excel")
    parser.add_argument("--source-dir", type=Path, default=DEFAULT_SOURCE_DIR, help="柏連正式文件資料夾")
    parser.add_argument("--main-list", type=Path, default=DEFAULT_MAIN_LIST_PATH, help="正式程序主清單 docx")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_PATH, help="輸出的 Excel 路徑")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    rows = build_master_rows(args.source_dir, args.main_list)
    write_workbook(rows, args.output, args.source_dir, args.main_list)
    print(args.output)


if __name__ == "__main__":
    main()
